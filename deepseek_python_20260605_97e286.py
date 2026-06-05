from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('JADESOLA CHATBOT – NAVIGATION SYSTEM FOR REDEEMER\'S UNIVERSITY', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# CHAPTER FOUR
doc.add_heading('CHAPTER FOUR', level=1)
doc.add_heading('SYSTEM IMPLEMENTATION AND RESULTS DISCUSSION', level=2)

doc.add_heading('Introduction', level=3)
doc.add_paragraph('This chapter describes the Chatbot implementation aimed to solve navigation problems in Redeemer’s University.')

doc.add_heading('Implementation Tools', level=3)
doc.add_paragraph('The tools utilized in the development of this software are described as follows:')

doc.add_heading('4.1.1 Google Dialogflow', level=4)
doc.add_paragraph('Designing and integrating a conversational user interface into your mobile app, online application, device, bot, interactive voice response system, etc. is simple using Google Dialogflow, a platform for natural language understanding. It is responsible for the integration of the Chatbot in the website.')

doc.add_heading('4.1.2 HyperText Markup Language (HTML)', level=4)
doc.add_paragraph('The HyperText Markup Language or HTML is the standard markup language for documents designed to be displayed in a web browser. HTML allows creating organized documents by expressing structural semantics for text elements like headings, paragraphs, links, lists, quotations and other elements.')

doc.add_heading('4.1.3 Cascading Style Sheet (CSS)', level=4)
doc.add_paragraph('Cascading Style Sheets (CSS) is a language for creating style sheets that describe how a document written in a markup language like HTML should be presented. The design, layout, and display options for the various devices and screen sizes that web pages are displayed on are defined using CSS. It can manage the design of several web pages at once. CSS files contain external stylesheets, which are saved as .css.')

doc.add_heading('4.1.4 JavaScript', level=4)
doc.add_paragraph('Along with HTML and CSS, the programming language JavaScript is one of the foundational elements of the World Wide Web. JavaScript enables interaction between programs, enabling interactivity within a website or application. Applications require an interaction engine to function. JavaScript was used to ensure interactivity in this Chatbot application.')

doc.add_heading('4.1.5 Bootstrap', level=4)
doc.add_paragraph('A free and open-source CSS framework called Bootstrap is designed for frontend web development that prioritizes mobile responsiveness. It can be downloaded and used at no cost. It includes design templates for typography, forms, buttons, navigation, and other interface elements in HTML, CSS, and (optionally) JavaScript. It makes creating responsive designs simpler. This tool was applied in this application.')

doc.add_heading('4.2 Development Environment', level=3)
doc.add_paragraph('The software and hardware environment in which this Chatbot was developed is described in Tables 4.1 and 4.2 respectively.')

doc.add_heading('Table 4.1: Software Environment', level=4)
p = doc.add_paragraph()
p.add_run('Operating System from the server: ').bold = True
p.add_run('Any Operating System\n')
p.add_run('Operating System for the client: ').bold = True
p.add_run('Windows OS, Android, Linux, Ubuntu\n')
p.add_run('Web Server: ').bold = True
p.add_run('GitHub\n')
p.add_run('Web Browser: ').bold = True
p.add_run('Google Chrome, Microsoft Edge, Safari, Internet Explorer, Mozilla Firefox, UC Browser, Opera\n')
p.add_run('Development Platform: ').bold = True
p.add_run('Google Dialogflow')

doc.add_heading('Table 4.2: Hardware Environment', level=4)
doc.add_paragraph('Processor: Intel® Core™ i3-2370M | Processor Speed: CPU @ 2.40GHz | Hard Disk Space: 512GB | RAM: 4GB')

doc.add_heading('Implementation Process', level=3)

doc.add_heading('4.3.1 Home page', level=4)
doc.add_paragraph('The Home page (Figure 4.1) is the landing page seen upon loading the website. It is made up of several images, icons and buttons. It includes a navigational bar which contains the Home, About, FAQ, Images and Contact. It also contains the Login button, which directs the user to the login page.')

doc.add_heading('The About page', level=4)
doc.add_paragraph('The About page (Figure 4.2) contains information regarding the Chatbot such as its aim as well as benefits.\n\n"Jadesola is a chatbot navigation system that helps new students, external workers and guests of Redeemer’s University or anyone unfamiliar to the environment navigate their way from one location to another on campus. Some benefits of Jadesola bot are: It is platform independent. Location. Navigation accessibility."')

doc.add_heading('FAQ page', level=4)
doc.add_paragraph('The FAQ page (Figures 4.3 and 4.3a) consists of some frequently asked questions by users.\n\n')
doc.add_paragraph('• Hostel to Lecture Room: Move straight down and make a LEFT turn at the PROPHET MOSES HALL or QUEEN ESTHER HALL signboard. Head downwards and make a LEFT turn at the Faculty of Law. Continue downwards and make the third RIGHT turn at Boja Theatres. Proceed downwards and make a LEFT turn onto LR 1-24.')
doc.add_paragraph('• School Gate to Guest House: Move straight down and make the second LEFT turn after the main public parking. Head down and make a RIGHT turn at the Faculty of Law. Continue downwards and make a LEFT turn onto the School House.')
doc.add_paragraph('• Hostel to Auditorium: Move down and make a LEFT turn at the QUEEN ESTHER HALL or PROPHET MOSES HALL signboard. Head straight down and make the third LEFT turn at the Faculty of Law. Continue downwards and make the second RIGHT turn at Peace Park. Move straight down and stop at Entrance 1 of the Auditorium.')
doc.add_paragraph('• School gate to Hostel: Head straight down and make the second LEFT turn after the main public parking. Head down and make a RIGHT turn at the Faculty of Law. Continue downwards and make the second RIGHT turn onto Queen Esther Hall.')

doc.add_heading('The Contact page', level=4)
doc.add_paragraph('The Contact page (Figures 4.5 and 4.5a) consists of information on how to get to or reach the developer of this project. It consists of phone number, email, and mapped location.\n\nADDRESS: P.M.B. 230, Ede, Osun State, Nigeria.\nPHONE NUMBER: +234 905 317 7566\nEMAIL: reddeemer@email.com\n\n© Copyright Redeemer’s University. All Rights Reserved. 2022')

doc.add_heading('The Login page', level=4)
doc.add_paragraph('The Login page (Figures 4.6 and 4.6a) contains fields for username and password. Users input these details, get validated and proceed to gain access to the bot. Users without accounts then click "Don’t have an account?" button and proceed to create one.\n\nSign In – To create an account, please sign in. Username, Password, Remember me, Forgot your password?')

doc.add_heading('The Images page', level=4)
doc.add_paragraph('Figure 4.4 & 4.4a: The Images Page\nSOME IMAGES OF REDEEMER’S UNIVERSITY – Some are a sample of the pictures of Redeemer’s University')

doc.add_heading('4.3.7 The Bot', level=4)
doc.add_paragraph('This is the chatroom (Figure 4.7) in which the user then converses with the bot whose alias is Jadesola. The user asks a question regarding the route from the location he or she is in at that point in time, to the location he or she desires to get to and receives a response from the bot.')

doc.add_heading('4.4 Typical Conversations with The Bot', level=3)

doc.add_heading('4.4.1 Salutations', level=4)
doc.add_paragraph('Jadesola can converse with a user at the salutation level as shown in Figure 4.8. When the user inputs "Hello" or "Hi", "Good morning" etc., Jadesola would reply; "Hi! How are you doing?"\n\nUser: Hello → Jadesola: Hi! How are you doing?\nUser: Who created you? → Jadesola: Onuigbo Ifeanyichukwu created me.')

doc.add_heading('Navigation Questions (Figures 4.10 to 4.19)', level=4)
doc.add_paragraph('• How to get to Boja from Girls Hostel? → Move down the road and make a LEFT turn at the junction by the QUEEN ESTHER HALL playground. Move down the road and make the LEFT turn at the junction. Move down the road and make the RIGHT turn at the junction by BOJA.')
doc.add_paragraph('• Girls Hostel to Chapel → Move straight down till you reach the ATM junction. Make a LEFT turn and move down till you get to the Faculty of Law. Make a LEFT turn and walk down till you get to Peace Park. Make a RIGHT turn straight down to Entrance 1 of the auditorium.')
doc.add_paragraph('• How can I get to Staff Quarters from the School Gate? → Proceed down the road and make a RIGHT turn by the family of Law. Then move downwards and make the fourth street turn on to Staff Quarters.')
doc.add_paragraph('• Boys Hostel to Bakery → Head straight down and make a LEFT turn at the PROPHET MOSES HALL signboard. Move downwards and make another LEFT turn at the family of Law. Continue downwards and make a LEFT turn on to the unlettered road and make a LEFT turn at the junction by the QUEEN ESTHER HALL signboard. Move down the road and make a LEFT turn at the junction by the Faculty of Law. Continue down the road and make a LEFT opposite the Unity Bank ATM on to the Business District (Container)')
doc.add_paragraph('• How do I get to Mimis from Guest House? → Start moving down the road and make a RIGHT turn at the junction by the PROPHET MOSES HALL signboard. Continue straight down on to the unlettered road and arrive at Mimis')
doc.add_paragraph('• Girls Hostel to Container → Move down the road and make a LEFT turn at the junction by the QUEEN ESTHER HALL signboard. Move down the road and make a LEFT turn at the junction by the Faculty of Law. Continue down the road and make a LEFT opposite the Unity Bank ATM on to the Business District (Container)')
doc.add_paragraph('• How do I locate Boys Hostel from School Gate? → I went downstairs and make the turn left to turn left to the Manna house. I found it. I found it down the road and make a right turn by the Faculty of Law. Then move towards and make the first right turn on to Boy\'s hostel.')
doc.add_paragraph('• How do I get to VIC\'s office from School Gate? → I went straight to the right down the road and make the turn left. I found it. I found it down the road and make a right turn by the Faculty of Law. Then move towards and make the first right turn on to VIC\'s office.')

doc.add_heading('Usability Test', level=3)
doc.add_paragraph('Usability Test was carried out to analyze how the system achieves the six (6) goals of usability: Memorability, Safety, Learnability, Efficiency, Effectiveness and Utility. Questions were asked in accordance to the goals and the table below (Table 4.3) contains the results of the test.')

doc.add_heading('Table 4.3 Results of Usability Test', level=4)
doc.add_paragraph('Memorability: 70% Strongly Agree, 30% Agree')
doc.add_paragraph('Safety: 60% Strongly Agree, 40% Agree')
doc.add_paragraph('Learnability: 80% Strongly Agree, 20% Agree')
doc.add_paragraph('Efficiency: 80% Strongly Agree, 20% Agree')
doc.add_paragraph('Effectiveness: 90% Strongly Agree, 10% Agree')
doc.add_paragraph('Utility: 70% Strongly Agree, 30% Agree')

doc.add_paragraph('In the question of memorability, 70% of students strongly agreed that they were able to remember how to start a conversation with the bot, 30% of students agreed that they were able to remember how to start a conversation with the bot.')
doc.add_paragraph('In the question of safety, 60% of students strongly agreed that they could delete a question in which mistakes were made, 40% of students agreed that they could delete a question in which mistakes were made.')
doc.add_paragraph('In the question of learnability, 80% of students strongly agreed that they could navigate actions on the interface of the website and bot, 20% of students agreed that they could navigate actions on the interface of the website and bot.')
doc.add_paragraph('In the question of efficiency, 80% of students strongly agreed that they could use the interface of the system on their own after learning it, 20% of students agreed that they could use the interface of the system on their own after learning it.')
doc.add_paragraph('In the question of effectiveness, 90% of students strongly agreed that they were able to access the information they needed from the bot, 10% agreed that they were able to access the information they needed from the bot.')
doc.add_paragraph('In the question of utility, 70% of students strongly agreed that the system provided all the functionalities they needed, 30% of students agreed that the system provided all the functionalities they needed.')

# CHAPTER FIVE
doc.add_page_break()
doc.add_heading('CHAPTER FIVE', level=1)
doc.add_heading('SUMMARY, CONCLUSION AND RECOMMENDATIONS', level=2)

doc.add_heading('Introduction', level=3)
doc.add_paragraph('This chapter presents the summary, conclusion, limitations recommendations and future works of the Chatbot developed in this project.')

doc.add_heading('Summary', level=3)
doc.add_paragraph('Chatbot system for campus navigation was created and tested for the sole purpose of solving navigational issues in Redeemer’s University, which occur majorly with individuals who are unfamiliar with the environment. A brief literature review was given to understand fundamental parts and aspects of this project as well as a review of related works. The system was designed adopting the waterfall model and using Unified Modelling Language (UML) diagrams. The primary tools used in the implementation of the Chatbot were Google Dialogflow, HTML, CSS, JavaScript and Bootstrap and a usability test was carried out on the system. The Chatbot developed in this project aimed to appropriately respond to users’ queries about how to navigate the campus environment. The users were asked means by which the Chatbot can be improved and the general feedback gotten included the addition of images, addition of emoticons and the introduction of a voice response feature.')

doc.add_heading('5.2 Conclusion', level=3)
doc.add_paragraph('The primary purpose of this project as stated in the summary was to solve navigational issues on campus, to avoid missing your way or to guide in looking for the directions from a starting point to the desired point. With the world being digitized, people spend more time on smart devices. This Chatbot aims to eliminate the need to ask passers-by for the directions to a desired location, and would instead help users solve this problem on smart devices.')

doc.add_heading('5.3 Limitation of the Study', level=3)
doc.add_paragraph('The following are the limitation of the study:')
doc.add_paragraph('1. It was difficult including questions involving every location on campus due to the large number of focal points on campus.', style='List Bullet')
doc.add_paragraph('2. The Chatbot cannot execute extended discussions with users, as it was programmed to respond to salutations and routing questions.', style='List Bullet')

doc.add_heading('5.4 Recommendations', level=3)
doc.add_paragraph('The Redeemer’s University is encouraged to adopt this Chatbot to enhance navigation on campus.')
doc.add_paragraph('The system can be improved in future by carrying out the following:')
doc.add_paragraph('1. Addition of images along with the response to give users an idea what each location looks like.', style='List Bullet')
doc.add_paragraph('2. Addition of emoticons.', style='List Bullet')
doc.add_paragraph('3. Addition of media and voice notes.', style='List Bullet')
doc.add_paragraph('4. Addition of more questions and responses to cover wider areas of the campus.', style='List Bullet')

# Save
doc.save('Jadesola_Chatbot_Report.docx')
print("✅ Report saved as 'Jadesola_Chatbot_Report.docx'")